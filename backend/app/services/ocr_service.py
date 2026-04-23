from azure.ai.formrecognizer.aio import DocumentAnalysisClient
from azure.identity.aio import ManagedIdentityCredential
from app.config import get_settings
import io

settings = get_settings()


class OCRService:
    def __init__(self):
        self.endpoint = settings.AZURE_DOCUMENT_INTELLIGENCE_ENDPOINT

    async def extract_text(self, file_bytes: bytes, content_type: str) -> dict:
        # ── DOCX — extract directly with python-docx ──────────────────────────
        if content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            return self._extract_docx(file_bytes)

        # ── TXT — read directly ───────────────────────────────────────────────
        if content_type == "text/plain":
            return self._extract_txt(file_bytes)

        # ── PDF + Images — Azure Document Intelligence ────────────────────────
        if settings.AZURE_DOCUMENT_INTELLIGENCE_KEY:
            from azure.core.credentials import AzureKeyCredential
            credential = AzureKeyCredential(settings.AZURE_DOCUMENT_INTELLIGENCE_KEY)
            client = DocumentAnalysisClient(
                endpoint=self.endpoint,
                credential=credential,
            )
            async with client:
                poller = await client.begin_analyze_document(
                    model_id="prebuilt-document",
                    document=file_bytes,
                )
                result = await poller.result()
        else:
            credential = ManagedIdentityCredential()
            async with credential:
                client = DocumentAnalysisClient(
                    endpoint=self.endpoint,
                    credential=credential,
                )
                async with client:
                    poller = await client.begin_analyze_document(
                        model_id="prebuilt-document",
                        document=file_bytes,
                    )
                    result = await poller.result()

        return self._parse_result(result)

    def _extract_docx(self, file_bytes: bytes) -> dict:
        """Extract text from DOCX using python-docx."""
        try:
            from docx import Document
            doc = Document(io.BytesIO(file_bytes))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            full_text = "\n".join(paragraphs)
            return {
                "full_text": full_text,
                "pages": [{"page_number": 1, "lines": paragraphs}],
                "page_count": 1,
                "tables": [],
                "key_value_pairs": [],
                "word_count": len(full_text.split()) if full_text else 0,
            }
        except Exception as e:
            return {
                "full_text": f"DOCX extraction failed: {str(e)}",
                "pages": [], "page_count": 0,
                "tables": [], "key_value_pairs": [], "word_count": 0,
            }

    def _extract_txt(self, file_bytes: bytes) -> dict:
        """Extract text from TXT file."""
        try:
            full_text = file_bytes.decode("utf-8", errors="ignore")
            lines = [l for l in full_text.splitlines() if l.strip()]
            return {
                "full_text": full_text,
                "pages": [{"page_number": 1, "lines": lines}],
                "page_count": 1,
                "tables": [],
                "key_value_pairs": [],
                "word_count": len(full_text.split()) if full_text else 0,
            }
        except Exception as e:
            return {
                "full_text": f"TXT extraction failed: {str(e)}",
                "pages": [], "page_count": 0,
                "tables": [], "key_value_pairs": [], "word_count": 0,
            }

    def _parse_result(self, result) -> dict:
        pages_text = []
        for page in result.pages:
            page_content = {
                "page_number": page.page_number,
                "width": page.width,
                "height": page.height,
                "lines": [line.content for line in page.lines] if page.lines else [],
            }
            pages_text.append(page_content)

        tables = []
        for table in result.tables:
            table_data = {
                "row_count": table.row_count,
                "column_count": table.column_count,
                "cells": [
                    {
                        "row_index": cell.row_index,
                        "column_index": cell.column_index,
                        "content": cell.content,
                    }
                    for cell in table.cells
                ],
            }
            tables.append(table_data)

        key_value_pairs = []
        if result.key_value_pairs:
            for kv in result.key_value_pairs:
                if kv.key and kv.value:
                    key_value_pairs.append({
                        "key": kv.key.content,
                        "value": kv.value.content,
                    })

        full_text = result.content if result.content else ""

        return {
            "full_text": full_text,
            "pages": pages_text,
            "page_count": len(result.pages),
            "tables": tables,
            "key_value_pairs": key_value_pairs,
            "word_count": len(full_text.split()) if full_text else 0,
        }